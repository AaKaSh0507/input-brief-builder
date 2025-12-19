from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from typing import Dict, Any
import os
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class ExportService:
    def __init__(self):
        self.export_dir = "/app/backend/exports"
        os.makedirs(self.export_dir, exist_ok=True)
    
    def export_to_pdf(self, brief_data: Dict[str, Any], filename: str = None) -> str:
        """Export brief to PDF format"""
        try:
            if not filename:
                filename = f"brief_{brief_data['id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            file_path = os.path.join(self.export_dir, filename)
            
            doc = SimpleDocTemplate(file_path, pagesize=letter,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=18)
            
            story = []
            styles = getSampleStyleSheet()
            
            # Title style
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor='#1a1a1a',
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            # Add title
            title = Paragraph(f"<b>{brief_data.get('title', 'Event Input Brief')}</b>", title_style)
            story.append(title)
            story.append(Spacer(1, 0.3 * inch))
            
            # Add metadata
            metadata_style = styles['Normal']
            story.append(Paragraph(f"<b>Status:</b> {brief_data.get('status', 'N/A')}", metadata_style))
            story.append(Paragraph(f"<b>Event Type:</b> {brief_data.get('event_type', 'N/A')}", metadata_style))
            story.append(Paragraph(f"<b>Version:</b> {brief_data.get('version', 1)}", metadata_style))
            story.append(Paragraph(f"<b>Created:</b> {brief_data.get('created_at', 'N/A')}", metadata_style))
            story.append(Spacer(1, 0.5 * inch))
            
            # Add sections
            heading_style = styles['Heading2']
            normal_style = styles['Normal']
            
            sections = brief_data.get('sections', [])
            for section in sections:
                # Section heading
                section_title = f"{section.get('section_number', '')}. {section.get('section_name', '')}"
                story.append(Paragraph(f"<b>{section_title}</b>", heading_style))
                story.append(Spacer(1, 0.2 * inch))
                
                # Section content
                content = section.get('content', {})
                for field, value in content.items():
                    if value:
                        story.append(Paragraph(f"<b>{field}:</b> {value}", normal_style))
                        story.append(Spacer(1, 0.1 * inch))
                
                story.append(Spacer(1, 0.3 * inch))
            
            # Build PDF
            doc.build(story)
            
            return file_path
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            raise
    
    def export_to_word(self, brief_data: Dict[str, Any], filename: str = None) -> str:
        """Export brief to Word format (DOCX)"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            if not filename:
                filename = f"brief_{brief_data['id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            file_path = os.path.join(self.export_dir, filename)
            
            doc = Document()
            
            # Add title
            title = doc.add_heading(brief_data.get('title', 'Event Input Brief'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f"Status: {brief_data.get('status', 'N/A')}")
            doc.add_paragraph(f"Event Type: {brief_data.get('event_type', 'N/A')}")
            doc.add_paragraph(f"Version: {brief_data.get('version', 1)}")
            doc.add_paragraph(f"Created: {brief_data.get('created_at', 'N/A')}")
            doc.add_paragraph()
            
            # Add sections
            sections = brief_data.get('sections', [])
            for section in sections:
                # Section heading
                section_title = f"{section.get('section_number', '')}. {section.get('section_name', '')}"
                doc.add_heading(section_title, level=2)
                
                # Section content
                content = section.get('content', {})
                for field, value in content.items():
                    if value:
                        p = doc.add_paragraph()
                        p.add_run(f"{field}: ").bold = True
                        p.add_run(str(value))
                
                doc.add_paragraph()
            
            # Save document
            doc.save(file_path)
            
            return file_path
        except Exception as e:
            logger.error(f"Error exporting to Word: {str(e)}")
            raise

# Singleton instance
export_service = ExportService()
